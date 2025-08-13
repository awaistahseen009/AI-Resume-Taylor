import os
import PyPDF2
import pdfplumber
from docx import Document
import re
from typing import Optional

class ResumeProcessor:
    """Service for processing and extracting text from resume files"""
    
    def __init__(self):
        self.supported_formats = ['pdf', 'docx', 'txt']
    
    def extract_text_from_file(self, file_path: str) -> str:
        """Extract text from uploaded resume file"""
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        file_extension = file_path.lower().split('.')[-1]
        
        if file_extension == 'pdf':
            return self._extract_from_pdf(file_path)
        elif file_extension == 'docx':
            return self._extract_from_docx(file_path)
        elif file_extension == 'txt':
            return self._extract_from_txt(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    def _extract_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file using multiple methods for better accuracy"""
        text = ""
        
        # Try pdfplumber first (better for complex layouts)
        try:
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"pdfplumber failed: {e}")
            
        # Fallback to PyPDF2 if pdfplumber fails or returns empty
        if not text.strip():
            try:
                with open(file_path, 'rb') as file:
                    pdf_reader = PyPDF2.PdfReader(file)
                    for page in pdf_reader.pages:
                        page_text = page.extract_text()
                        if page_text:
                            text += page_text + "\n"
            except Exception as e:
                print(f"PyPDF2 failed: {e}")
                raise ValueError("Could not extract text from PDF file")
        
        return self._clean_text(text)
    
    def _extract_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return self._clean_text(text)
            
        except Exception as e:
            raise ValueError(f"Could not extract text from DOCX file: {str(e)}")
    
    def _extract_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                text = file.read()
            return self._clean_text(text)
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    text = file.read()
                return self._clean_text(text)
            except Exception as e:
                raise ValueError(f"Could not read text file: {str(e)}")
    
    def _clean_text(self, text: str) -> str:
        """Clean and normalize extracted text"""
        if not text:
            return ""
        
        # Remove excessive whitespace
        text = re.sub(r'\n\s*\n', '\n\n', text)
        text = re.sub(r' +', ' ', text)
        
        # Remove special characters that might cause issues
        text = re.sub(r'[^\w\s\-\.\,\(\)\[\]\@\#\$\%\&\*\+\=\:\;\!\?\'\"\n]', '', text)
        
        # Normalize line breaks
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
    
    def validate_resume_content(self, text: str) -> dict:
        """Validate if the extracted text looks like a resume"""
        validation = {
            'is_valid': True,
            'warnings': [],
            'suggestions': []
        }
        
        text_lower = text.lower()
        
        # Check minimum length
        if len(text.split()) < 50:
            validation['warnings'].append("Resume appears to be very short")
        
        # Check for common resume sections
        resume_sections = ['experience', 'education', 'skills', 'work', 'employment', 'projects']
        sections_found = sum(1 for section in resume_sections if section in text_lower)
        
        if sections_found < 2:
            validation['warnings'].append("Resume may be missing common sections (Experience, Education, Skills)")
        
        # Check for contact information
        has_email = '@' in text and '.' in text
        has_phone = bool(re.search(r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b', text))
        
        if not has_email:
            validation['suggestions'].append("Consider adding email address")
        if not has_phone:
            validation['suggestions'].append("Consider adding phone number")
        
        # Check for dates (employment history)
        date_patterns = [
            r'\b\d{4}\b',  # Year
            r'\b\d{1,2}/\d{4}\b',  # Month/Year
            r'\b(Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{4}\b'  # Month Year
        ]
        
        has_dates = any(re.search(pattern, text) for pattern in date_patterns)
        if not has_dates:
            validation['suggestions'].append("Consider adding employment dates")
        
        return validation
    
    def extract_contact_info(self, text: str) -> dict:
        """Extract contact information from resume text"""
        contact_info = {}
        
        # Extract email
        email_match = re.search(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        if email_match:
            contact_info['email'] = email_match.group()
        
        # Extract phone number
        phone_patterns = [
            r'\b\d{3}[-.\s]?\d{3}[-.\s]?\d{4}\b',
            r'\(\d{3}\)\s*\d{3}[-.\s]?\d{4}',
            r'\+\d{1,3}[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{4}'
        ]
        
        for pattern in phone_patterns:
            phone_match = re.search(pattern, text)
            if phone_match:
                contact_info['phone'] = phone_match.group()
                break
        
        # Extract LinkedIn profile
        linkedin_match = re.search(r'linkedin\.com/in/[\w-]+', text.lower())
        if linkedin_match:
            contact_info['linkedin'] = linkedin_match.group()
        
        # Extract GitHub profile
        github_match = re.search(r'github\.com/[\w-]+', text.lower())
        if github_match:
            contact_info['github'] = github_match.group()
        
        return contact_info
